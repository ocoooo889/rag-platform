"""
多格式文档文本提取 — 统一成纯文本后再走现有切片 / 向量化。

支持：.md /.markdown /.txt /.pdf /.docx /.html /.htm /.csv
PDF / Word 需可选依赖；缺失时给出明确错误提示。
"""

from __future__ import annotations

import csv
import html
import io
import logging
import re
from html.parser import HTMLParser
from pathlib import Path

logger = logging.getLogger(__name__)

# 知识库上传允许的后缀（小写，不含点）
ALLOWED_DOC_EXTENSIONS: frozenset[str] = frozenset(
    {
        "md",
        "markdown",
        "txt",
        "pdf",
        "docx",
        "html",
        "htm",
        "csv",
    }
)

_TEXT_EXTS = frozenset({"md", "markdown", "txt"})
_HTML_EXTS = frozenset({"html", "htm"})


class _HTMLTextExtractor(HTMLParser):
    """粗粒度去标签，保留可读正文。"""

    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: ANN001
        if tag.lower() in {"script", "style", "noscript"}:
            self._skip += 1
        elif tag.lower() in {"p", "div", "br", "li", "h1", "h2", "h3", "h4", "tr"}:
            self._chunks.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript"} and self._skip:
            self._skip -= 1
        elif tag.lower() in {"p", "div", "li", "h1", "h2", "h3", "h4", "tr"}:
            self._chunks.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        text = data.strip()
        if text:
            self._chunks.append(text)

    def get_text(self) -> str:
        raw = " ".join(self._chunks)
        raw = html.unescape(raw)
        raw = re.sub(r"[ \t]+\n", "\n", raw)
        raw = re.sub(r"\n{3,}", "\n\n", raw)
        raw = re.sub(r"[ \t]{2,}", " ", raw)
        return raw.strip()


def normalize_ext(filename: str | None) -> str:
    name = (filename or "").strip().lower()
    if "." not in name:
        return ""
    return name.rsplit(".", 1)[-1]


def is_allowed_doc_filename(filename: str | None) -> bool:
    return normalize_ext(filename) in ALLOWED_DOC_EXTENSIONS


def allowed_ext_label() -> str:
    """给人看的后缀列表。"""
    order = ["md", "txt", "pdf", "docx", "html", "htm", "csv", "markdown"]
    shown = [e for e in order if e in ALLOWED_DOC_EXTENSIONS]
    return "、".join(f".{e}" for e in shown)


def _read_plain_text(file_path: str) -> str:
    """UTF-8 优先，失败再试常见中文编码。"""
    last_err: Exception | None = None
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312"):
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError as e:
            last_err = e
            continue
    raise ValueError(f"无法解码文档，已尝试 utf-8/gbk：{last_err}")


def _extract_pdf(file_path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "未安装 pypdf，无法解析 PDF。请在 backend 目录执行: pip install pypdf"
        ) from e

    reader = PdfReader(file_path)
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            logger.warning("PDF 第 %s 页提取失败: %s", i + 1, exc)
            continue
        text = text.strip()
        if text:
            parts.append(text)
    content = "\n\n".join(parts).strip()
    if not content:
        raise ValueError("PDF 未能提取到文本（可能是扫描件/纯图片，需 OCR）")
    return content


def _extract_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "未安装 python-docx，无法解析 Word。请执行: pip install python-docx"
        ) from e

    doc = Document(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    # 表格单元格一并纳入，避免制度表被丢掉
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    content = "\n".join(parts).strip()
    if not content:
        raise ValueError("Word 文档无有效文本内容")
    return content


def _extract_html(file_path: str) -> str:
    raw = _read_plain_text(file_path)
    parser = _HTMLTextExtractor()
    parser.feed(raw)
    content = parser.get_text()
    if not content:
        raise ValueError("HTML 未能提取到有效正文")
    return content


def _extract_csv(file_path: str) -> str:
    """CSV 转成逐行可读文本，便于切片检索。"""
    raw = _read_plain_text(file_path)
    reader = csv.reader(io.StringIO(raw))
    lines: list[str] = []
    for row in reader:
        cells = [str(c).strip() for c in row if str(c).strip()]
        if cells:
            lines.append(" | ".join(cells))
    content = "\n".join(lines).strip()
    if not content:
        raise ValueError("CSV 无有效内容")
    return content


def load_document_text(file_path: str, filename: str | None = None) -> str:
    """
    按扩展名提取纯文本。
    filename 优先用于判型（上传存储名可能带 doc_id 前缀）。
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    # 优先用原始文件名后缀；否则用磁盘路径
    ext = normalize_ext(filename) or normalize_ext(path.name)
    if not ext:
        raise ValueError("无法识别文件类型（缺少扩展名）")
    if ext not in ALLOWED_DOC_EXTENSIONS:
        raise ValueError(f"不支持的文件类型 .{ext}，仅支持 {allowed_ext_label()}")

    if ext in _TEXT_EXTS:
        return _read_plain_text(str(path))
    if ext == "pdf":
        return _extract_pdf(str(path))
    if ext == "docx":
        return _extract_docx(str(path))
    if ext in _HTML_EXTS:
        return _extract_html(str(path))
    if ext == "csv":
        return _extract_csv(str(path))

    raise ValueError(f"未实现的文件类型: .{ext}")
